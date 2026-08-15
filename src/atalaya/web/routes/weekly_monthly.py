"""Vistas Semanal (§6.2) y Mensual (§6.3) + export DOCX de la síntesis."""
from __future__ import annotations

import io

from fastapi import APIRouter, Depends, Request, Response
from fastapi.responses import HTMLResponse
from sqlalchemy import desc, select
from sqlalchemy.orm import Session

from atalaya.config import THEME_LABELS_ES, WEEKLY_THEMES, load_countries
from atalaya.db.models import Article, MonthlySynthesis, WeeklyItem
from atalaya.web.deps import current_user, get_db, render

router = APIRouter()


@router.get("/weekly", response_class=HTMLResponse)
def weekly(request: Request, country: str | None = None, week: str | None = None,
           user_sess=Depends(current_user), db: Session = Depends(get_db)):
    user, sess = user_sess
    countries = {c: v for c, v in load_countries().items() if v.weekly}
    followed = [c for c in (user.countries or []) if c in countries]
    country = country or (followed[0] if followed else next(iter(countries)))

    weeks = [w for (w,) in db.execute(
        select(WeeklyItem.iso_week).distinct().order_by(desc(WeeklyItem.iso_week))).all()]
    week = week or (weeks[0] if weeks else None)

    by_theme: dict[str, list] = {th: [] for th in WEEKLY_THEMES}
    if week:
        rows = db.execute(
            select(WeeklyItem, Article).join(Article, WeeklyItem.article_id == Article.id)
            .where(WeeklyItem.country == country, WeeklyItem.iso_week == week)
            .order_by(Article.published_at.desc())
        ).all()
        for w, a in rows:
            by_theme.setdefault(w.theme, []).append({
                "url": a.url, "title": a.title, "summary": w.mini_summary_es,
                "date": a.published_at.date().isoformat() if a.published_at else None,
                "source": a.source_name,
            })
    return render(request, "weekly.html", user=user, csrf=sess.csrf_token,
                  weekly_countries=countries, sel_country=country,
                  weeks=weeks, sel_week=week, by_theme=by_theme,
                  theme_labels=THEME_LABELS_ES, themes=WEEKLY_THEMES)


@router.get("/monthly", response_class=HTMLResponse)
def monthly(request: Request, country: str | None = None, month: str | None = None,
            user_sess=Depends(current_user), db: Session = Depends(get_db)):
    user, sess = user_sess
    countries = {c: v for c, v in load_countries().items() if v.weekly}
    followed = [c for c in (user.countries or []) if c in countries]
    country = country or (followed[0] if followed else next(iter(countries)))

    months = [m for (m,) in db.execute(
        select(MonthlySynthesis.month).where(MonthlySynthesis.country == country)
        .distinct().order_by(desc(MonthlySynthesis.month))).all()]
    month = month or (months[0] if months else None)
    synthesis = db.scalar(select(MonthlySynthesis).where(
        MonthlySynthesis.country == country, MonthlySynthesis.month == month)) if month else None

    return render(request, "monthly.html", user=user, csrf=sess.csrf_token,
                  weekly_countries=countries, sel_country=country,
                  months=months, sel_month=month, synthesis=synthesis,
                  theme_labels=THEME_LABELS_ES, themes=WEEKLY_THEMES)


@router.get("/monthly/export.docx")
def monthly_docx(country: str, month: str,
                 user_sess=Depends(current_user), db: Session = Depends(get_db)):
    from docx import Document
    synthesis = db.scalar(select(MonthlySynthesis).where(
        MonthlySynthesis.country == country, MonthlySynthesis.month == month))
    if not synthesis:
        return Response(status_code=404)
    c = load_countries().get(country)
    doc = Document()
    doc.add_heading(f"Síntesis mensual — {c.name if c else country} — {month}", level=0)
    doc.add_heading("Síntesis global del país", level=1)
    doc.add_paragraph(synthesis.overview_es or "")
    for theme in WEEKLY_THEMES:
        section = (synthesis.sections or {}).get(theme)
        if not section:
            continue
        doc.add_heading(THEME_LABELS_ES[theme], level=1)
        doc.add_paragraph(section.get("sintesis") or "")
        for art in section.get("articulos", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{art.get('fecha') or ''} — {art.get('titulo')} ({art.get('fuente') or ''}) — {art.get('url')}")
    doc.add_heading("Tabla de incidentes", level=1)
    incidents = synthesis.incidents or []
    if incidents:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        for i, head in enumerate(["Fecha", "Localización", "Nivel", "Categoría",
                                  "Descripción", "Fuente(s)"]):
            table.rows[0].cells[i].text = head
        for inc in incidents:
            row = table.add_row().cells
            row[0].text = inc.get("fecha") or ""
            row[1].text = inc.get("localizacion") or ""
            row[2].text = inc.get("nivel") or ""
            row[3].text = inc.get("categoria") or ""
            row[4].text = inc.get("descripcion") or ""
            row[5].text = "; ".join(s.get("url", "") for s in inc.get("fuentes", []))
    else:
        doc.add_paragraph("Sin incidentes registrados en el período.")
    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=sintesis-{country}-{month}.docx"},
    )
